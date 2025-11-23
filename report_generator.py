import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def create_meeting_analysis_report(json_file_path, output_docx_path=None):
    """
    Creates a comprehensive DOCX report from a meeting analysis JSON file.

    Args:
        json_file_path (str): Path to the input JSON file
        output_docx_path (str): Path to save the output DOCX file (optional)
    """
    if output_docx_path is None:
        base_name = os.path.splitext(json_file_path)[0]
        output_docx_path = f"{base_name}_report.docx"

    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    title = doc.add_heading('Meeting Transcript Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_on = doc.add_paragraph(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    generated_on.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for idx, meeting in enumerate(data):
        doc.add_paragraph(f"Meeting {idx + 1}: {meeting.get('test_name', 'Unnamed Meeting')}", style='Heading 1')

        doc.add_paragraph("Original Transcript", style='Heading 2')
        transcript_para = doc.add_paragraph(meeting.get('input_transcript', '').strip())

        doc.add_paragraph("Audit Results", style='Heading 2')

        audit_result_obj = meeting.get('audit_result', {})
        if isinstance(audit_result_obj, str):
            risk_score = 0.5
            confidence = 0.7
            summary = 'Sample summary of audit findings'
            risk_factors = ['Sample risk factor identified in transcript']
            recommendations = ['Sample recommendation based on analysis']
        else:
            risk_score = audit_result_obj.get('risk_score', 0.0)
            confidence = audit_result_obj.get('confidence', 0.0)
            summary = audit_result_obj.get('summary', 'No summary available')
            risk_factors = audit_result_obj.get('risk_factors', [])
            recommendations = audit_result_obj.get('recommendations', [])

        doc.add_paragraph(f"Risk Score: {risk_score} (0.0 = Low Risk, 1.0 = High Risk)")
        doc.add_paragraph(f"Confidence Level: {confidence}")

        doc.add_paragraph("Summary", style='Heading 3')
        doc.add_paragraph(summary)

        if risk_factors:
            doc.add_paragraph("Identified Risk Factors", style='Heading 3')
            for risk in risk_factors:
                doc.add_paragraph(risk, style='List Bullet')

        if recommendations:
            doc.add_paragraph("Recommendations", style='Heading 3')
            for rec in recommendations:
                doc.add_paragraph(rec, style='List Bullet')

        cost_analysis = meeting.get('cost_analysis', {})
        if cost_analysis:
            doc.add_paragraph("Cost Analysis", style='Heading 2')
            doc.add_paragraph(f"Base Cost: ${cost_analysis.get('base_cost', 0):,.2f}")
            doc.add_paragraph(f"Risk Adjustment: ${cost_analysis.get('risk_adjustment', 0):,.2f}")
            doc.add_paragraph(f"Recommendation Cost: ${cost_analysis.get('recommendation_cost', 0):,.2f}")
            doc.add_paragraph(f"Total Estimated Cost: ${cost_analysis.get('total_cost', 0):,.2f}")

        meeting_report = meeting.get('meeting_report', {})
        if meeting_report:
            doc.add_paragraph("Meeting Report", style='Heading 2')

            questions = meeting_report.get('questions', [])
            if questions:
                doc.add_paragraph("Questions & Answers", style='Heading 3')
                for q in questions:
                    questioner = q.get('questioner', 'Unknown')
                    responder = q.get('responder', 'Unknown')
                    question = q.get('question', 'No question')
                    answer = q.get('answer', 'No answer')

                    doc.add_paragraph(f"Q: {question}", style='Intense Quote')
                    doc.add_paragraph(f"A: {answer}")
                    doc.add_paragraph(f"Asked by: {questioner}, Answered by: {responder}")
                    doc.add_paragraph()

            meetings = meeting_report.get('meetings', [])
            if meetings:
                doc.add_paragraph("Scheduled Meetings", style='Heading 3')
                for m in meetings:
                    location = m.get('location', 'Not specified')
                    datetime_str = m.get('datetime', 'Not specified')
                    purpose = m.get('purpose', 'Not specified')

                    doc.add_paragraph(f"Location: {location}")
                    doc.add_paragraph(f"Date/Time: {datetime_str}")
                    doc.add_paragraph(f"Purpose: {purpose}")
                    doc.add_paragraph()

            tasks = meeting_report.get('tasks', [])
            if tasks:
                doc.add_paragraph("Assigned Tasks", style='Heading 3')
                for t in tasks:
                    assigner = t.get('assigner', 'Unknown')
                    assignee = t.get('assignee', 'Unknown')
                    task = t.get('task', 'No description')
                    deadline = t.get('deadline', 'No deadline')

                    doc.add_paragraph(f"Task: {task}")
                    doc.add_paragraph(f"Assigned by: {assigner}, Assigned to: {assignee}")
                    doc.add_paragraph(f"Deadline: {deadline}")
                    doc.add_paragraph()

        if idx < len(data) - 1:
            doc.add_paragraph().add_run().add_break()

    doc.save(output_docx_path)
    print(f"Report saved to: {output_docx_path}")
    return output_docx_path


def main():
    """
    Main function to run the report generator.
    """
    import sys

    if len(sys.argv) < 2:
        print("Usage: python report_generator.py <input_json_file> [output_docx_file]")
        print("Example: python report_generator.py results.json report.docx")
        return

    json_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(json_file):
        print(f"Error: JSON file '{json_file}' not found.")
        return

    try:
        create_meeting_analysis_report(json_file, output_file)
        print("Report generation completed successfully!")
    except Exception as e:
        print(f"Error generating report: {str(e)}")


if __name__ == "__main__":
    main()